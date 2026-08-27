"""Renders the generated Markdown notes into a .docx that mirrors the
configured template's section structure as closely as practical."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

import re

def _add_markdown_runs(paragraph, text: str):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)

def write_docx(meeting_dir: Path, meeting_title: str, meeting_date: str, notes_markdown: str) -> Path:
    doc = Document()

    doc.add_heading(meeting_title, level=0)
    doc.add_paragraph(meeting_date)

    for raw_line in notes_markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_markdown_runs(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            _add_markdown_runs(p, line)
            p.paragraph_format.space_after = Pt(6)

    path = meeting_dir / "notes.docx"
    doc.save(str(path))
    return path
